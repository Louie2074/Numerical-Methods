
import io
import matplotlib.pyplot as plt
import numpy as np
from scipy.interpolate import make_interp_spline, BSpline
import docx
from cmath import cos, pi, sin, tan
import numpy as np
import math
from scipy import integrate

# library allows for microsoft docx export
mydoc = docx.Document("./Doc1.docx")


'''Question 1A of newtons method'''


def q1A():
    def function(x): return ((sin(x)/(x-1)**2)+5 *
                             (cos(0.5*(x**2))**2)+3*x+1.5).real
    def Derivative(x): return (-10*x*cos(x**2/2)*sin(x**2/2) -
                               ((2*sin(x))/(x-1)**3)+((cos(x))/(x-1)**2)+3).real
    approx = newton(-1, function, Derivative, 1e-6, True, 100)


'''Question 1B of newtons method'''


def q1B():
    def function(xy):
        x, y = xy
        return [2*x*sin(1.5*x)-y,
                4*(x-6.7)**5-y-2]

    def Derivative(xy):
        x, y = xy
        return [[3*x*cos(1.5*x)+2*sin(1.5*x), -1],
                [20*(x-6.7)**4, -1]]
    approx = newton([1.0, -2.0], function,
                    Derivative, 1e-6, True, 100,)
    tab = toTable(approx)
    print(tab)
    mydoc.add_paragraph(tab)
    mydoc.save("./Doc1.docx")


'''Question 1C of newtons method'''


def q1C():
    def function(xyz):
        x, y, z = xyz
        return [3*x-2*y-z-4,
                -x+3*y-z-8,
                x-z-2]

    def Derivative(xyz):
        x, y, z = xyz
        return [[3, -2, -1],
                [-1, 3, -1],
                [1, 0, -1]]

    approx = newton([1.0, 1.0, 1.0], function,
                    Derivative, 1e-6, False, 100)


'''Question 1D of newtons method'''


def q1D():
    def function(x): return ((2*(x**4))-(6*(x**3))-(8*(x**2))+(24*x)).real
    def Derivative(x): return ((8*(x**3))-(18*(x**2))-(16*x)+24).real
    approx1 = newton(-1.5, function, Derivative, 1e-6, True, 100)
    approx2 = newton(1.5, function, Derivative, 1e-6, True, 100)
    approx3 = newton(2.7, function, Derivative, 1e-6, True, 100)
    approx4 = newton(-0.5, function, Derivative, 1e-6, True, 100)


'''Question 2A of Eulers method'''


def q2A():
    def function(x, y): return (-4*sin((pi/6)*math.log(y))*(x-1)).real
    approx = euler(0, 4, 2, 0.1, function, True, None, None)
    plotG(approx, 0.1, "Eulers", False, None)


'''Question 2B of Eulers method'''


def q2B():
    def function(x, y, z): return -sin(z*(2*x+y))
    def function2(x, y): return math.e**(-0.05*x)*cos(4*y)
    euler(0, 7, 0, 1e-5, function, False, function2, -1)


'''Question 2C of Eulers method'''


def q2C():
    def function(x, y): return (y+x**2-math.e**(1.5*y)).real
    approx1 = euler(0, 10, 0, 1e-5, function, True, None, None)
    approx2 = euler(0, 10, 0, 0.05, function, True, None, None)
    plotG(approx1, 1e-5, "Eulers", False, None)
    plotG(approx2, 0.05, "Eulers", False, None)


'''Question 3A of Midpoint method'''


def q3A():
    def function(x): return tan(x-(math.pi/3))+4*cos(x**3)
    approx = midpoint(0, 2.5, 1e-5, function)
    plotG(approx, 1e-5, "Midpoint", True, 2.5)


'''Question 3B of Midpoint method'''


def q3B():
    def function(x): return sin(2*x)+cos(3*x)
    approx1 = midpoint(0, 4*math.pi, 1e-5, function)
    approx2 = midpoint(0, 4*math.pi, 1, function)
    plotG(approx1, 1e-5, "Midpoint", True, 4*math.pi)
    plotG(approx2, 1, "Midpoint", True, 4*math.pi)


def newton(x0, f, d, epsilon, display=True, max_iter=100):
    """Newtons Method for solving algebraic equations

    Args:
        x0: Initial Guess
        f: function to be solved
        d: derivative of function to be solved
        epsilon: Accuracy threshold of answer
        display (bool, optional): Optional argument if user wants details of calculation. Defaults to True.
        max_iter (int, optional): Optional argument if user wants to alter the amount of iterations. Defaults to 100.

    Returns:
        coords: x and y values to be tabulated or graphed
    """

    xn = x0
    coords = {}

    # Non jacobian Newtons method
    if not isinstance(x0, list):
        for i in range(0, max_iter):
            fxn = f(xn)
            coords[i] = [xn.real, fxn.real]
            # If user wants the details then plot
            if display:
                plotNewton(f, d, xn, i+1)
            # If the answer is accurate enough then stop
            if abs(fxn) < epsilon:
                print(
                    f'Initial Guess: {x0}, Found solution after {i+1} iterations. The answer is {xn}')
                return coords if display else None
            Dfxn = d(xn)
            # If the derivitive is 0 then quit
            if Dfxn == 0:
                print('Derivative is 0. No solution found.')
                return coords if display else None
            xn = xn - (fxn*(Dfxn**-1))

    # Jacobian Newtons method
    else:
        for i in range(max_iter):

            Jacobian = np.array(d(xn))
            NonJacobian = np.array(f(xn))

            diff = np.linalg.solve(Jacobian, -NonJacobian)
            coords[i] = [xn[0].real, xn[1].real]
            xn = xn + diff

            # If the answer is accurate enough then stop
            if np.linalg.norm(diff) < epsilon:
                print(f'{[i.real for i in xn]}')
                return coords if display else None

    # If max iterations have been reached and no answer was found
    print('No solution found.')
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def euler(x0, n, y0, h, f, display=True, f2=None, z0=None):
    """Eulers method for solving differential equations

    Args:
        x0: Starting X value
        n: End X value
        y0: Initial starting condition
        h: Step size
        f: function to be solved
        display (bool, optional): Optional argument if user wants details of calculation. Defaults to True.
        f2 (optional): Optional function argument if using to solve system of equation. Defaults to None.
        z0 (optional): Optional function starting condition argument if using to solve system of equations. Defaults to None.

    Returns:
        coords: x and y values to be tabulated or graphed
    """
    yn = y0
    zn = z0
    xn = x0
    coords = {}
    count = 0

    try:
        for i in np.arange(float(x0), float(n), h):
            # Not a system of equations
            if(f2 == None):
                coords[i] = yn.real
                new = yn+(h*f(i, yn))
                count += 1
                yn = new
            # Used if input is a system
            else:
                yn = yn+(h*f(i, yn, zn))
                zn = zn+(h*f2(i, yn))
                xn = i
                count += 1
    # Throws an error message if there is an error in the calculations
    except:
        print("An error occurred in the calculations, please check the inputs")

    # if input is not a system then output the y value, else output xyz values
    if(f2 == None):
        print(
            f'Initial Guess: {x0}, Found solution after {count} iterations. The answer is {yn}')
    else:
        print(
            f'Initial Guess: {x0}, Found solution after {count} iterations. The answer is {xn.real,yn.real,zn.real}')
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def midpoint(start, stop, step, f, display=True):
    """Midpoint rule for integrating an equation

    Args:
        start: Starting X value
        stop: End X value
        step: Step size
        f: function to be solved
        display (bool, optional): Optional argument if user wants details of calculation. Defaults to True.

    Returns:
         coords: x and y values to be tabulated or graphed
    """
    integral = 0
    coords = {}
    try:
        for i in np.arange(float(start), float(stop), step):
            coords[i] = f(i).real
            integral += f(((i+step)+i)/2)*step
        print("Integral is equal to: ", integral.real)
    # Throws an error message if there is an error in the calculations
    except:
        print("An error occurred in the calculations, please check the inputs")
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def midpointRungeKutta(x0, n, y0, h, f, display=True):
    """Midpoint Method for solving differential equations

    Args:
        x0: Starting X Value
        n: Ending X Value
        y0: Initial Starting Condition
        h: Step Size
        f: function to be solved
        display (bool, optional): Optional argument if user wants details of calculation. Defaults to True.

    Returns:
       coords: x and y values to be tabulated or graphed
    """
    yn = y0
    count = 0
    coords = {}
    try:
        for i in np.arange(float(x0), float(n), h):
            coords[i] = yn.real
            k1 = f(i, yn)
            k2 = f(i+(h/2), yn+(k1*h/2))
            count += 1
            yn = yn+(k2*h)

        print(
            f'Initial Guess: {x0}, Found solution after {count} iterations. The answer is {yn}')

    # Throws an error message if there is an error in the calculations
    except:
        print("An error occurred in the calculations, please check the inputs")
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def trapezoid(start, stop, step, f, display=True):
    """Trapezoid method for integrating an equation

    Args:
        start: Starting X Value
        stop: Ending X Value
        step: Step size
        f: function to be passed
        display (bool, optional): Optional argument if user wants details of calculation. Defaults to True.

    Returns:
        coords: x and y values to be tabulated or graphed
    """
    integral = 0
    coords = {}
    try:
        for i in np.arange(float(start), float(stop), step):
            coords[i] = f(i)
            integral += ((f(i)+f(i+step))/2)*step
        print("Integral is equal to: ", integral.real)
    # Throws an error message if there is an error in the calculations
    except:
        print("An error occurred in the calculations, please check the inputs")
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def rightpoint(start, stop, step, f, display=True):
    """Rightpoint method for integrating an equation

    Args:
        start: Starting X Value
        stop: Ending X Value
        step: Step Size
        f: function to be solved
        display (bool, optional):  Optional argument if user wants details of calculation. Defaults to True.

    Returns:
        coords: x and y values to be tabulated or graphed
    """
    integral = 0
    coords = {}
    try:
        for i in np.arange(float(start), float(stop), step):
            coords[i] = f(i)
            integral += f(i)*step
        print("Integral is equal to: ", integral.real)
    # Throws an error message if there is an error in the calculations
    except:
        print("An error occurred in the calculations, please check the inputs")
    # Returns the values to be graphed if the user specifies else returns none
    return coords if display else None


def toTable(coords):
    """Converts coordinate values into a tabular form

    Args:
        coords (_type_): Dictionary of coordinate values

    Returns:
        str: Table of values in a string
    """

    str = ""
    str += " "
    str += "{:<8} {:<15} {:<10}".format('Iteration', 'X', 'Y')
    for k, v in coords.items():
        x, y = v
        str += "\n"
        str += "{:<8} {:<15} {:<10}".format(k, x, y)
    str += " "
    return str


def plotG(dic, step, name, fill, max=None):
    """Plots end graph using the return values of calculations

    Args:
        dic (_type_): Dictionary of coordinate values
        step (_type_): Step Size
        name (_type_): Name of function
        fill (_type_): User option to fill areas under graph
        max (_type_, optional): Max x value of axis. Defaults to None.
    """

    # Sorts the x and y values of coordinates
    myList = dic.items()
    myList = sorted(myList)
    x, y = zip(*myList)

    # Plots the values
    plt.plot(x, y)
    plt.title(f"{name} method with step size of {step}")
    plt.xlabel("X")
    plt.ylabel("Y")

    # If the fill option is True then fills the area under the curve and plots x axis
    if fill:
        plt.fill_between(x, y, color='green', alpha=0.5)
        plt.hlines(0, 0, max, color='red')

    # Exports the graph into a microsoft word document
    figure = io.BytesIO()
    plt.savefig(figure, format='png')
    figure.seek(0)
    mydoc.add_picture(figure)
    mydoc.save('./Doc1.docx')
    plt.close()


def plotNewton(f, d, xn, iter):
    """Graphing method for newtons method that plots the function after every iteration

    Args:
        f: function to be graphed
        d: derivative of function to be graphed
        xn: Initial Guess
        iter: current iteration
    """
    # Sets up x and y values to be graphed and smooths the values out
    x = np.arange(-5, 5, 0.5)
    xnew = np.linspace(x.min(), x.max(), 300)
    y = []
    for i in x:
        try:
            y.append(f(i))
        except:
            y.append(10)
    spl = make_interp_spline(x, y, k=3)
    power_smooth = spl(xnew)

    # Sets up figure and axes of the graph
    fig = plt.figure(1)
    ax = fig.add_subplot(111)
    ax.spines['left'].set_position('zero')
    ax.spines['right'].set_color('none')
    ax.spines['bottom'].set_position('zero')
    ax.spines['top'].set_color('none')
    ax.xaxis.set_ticks_position('bottom')
    ax.yaxis.set_ticks_position('left')
    ax.set_xbound(-8, 8)
    ax.set_ybound(-8, 8)
    # plots the function
    ax.plot(xnew, power_smooth, linewidth=3, color='red', label="F(x)")
    # plots the tangent line
    ax.axline((xn, f(xn)), slope=d(xn), color='C0', label='Tangent')
    # plots the location of current guess
    ax.plot(xn, f(xn), marker="o", color="green",
            label=f"Guess \n X: {f'{xn:.2f}'} \n Y: {f'{f(xn):.2f}'}")

    # Title and legend of graph
    plt.title(f"Current Iteration: {iter}")
    plt.legend(loc="upper left")

    # Exports the graph into a microsoft word document
    figure = io.BytesIO()
    plt.savefig(figure, format='png')
    figure.seek(0)
    mydoc.add_picture(figure)
    mydoc.save('./Doc1.docx')
    plt.close()
